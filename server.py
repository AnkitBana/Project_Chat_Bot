"""
FastAPI backend for the AI Agent web interface.
Run with: python server.py
"""
import sys
import os
import asyncio
import json
import secrets
import datetime
sys.stdout.reconfigure(encoding="utf-8")

from fastapi import FastAPI, HTTPException, Depends, UploadFile, File
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
from typing import Optional, List
import uvicorn

# ── Auth helpers ──────────────────────────────────────────────────────────────
import bcrypt as _bcrypt_lib
from jose import JWTError, jwt

_SECRET_KEY   = os.environ.get("SECRET_KEY", secrets.token_hex(32))
_ALGORITHM    = "HS256"
_TOKEN_EXPIRE = 60 * 24  # minutes — 24 hours

_bearer = HTTPBearer(auto_error=False)

def _hash_password(password: str) -> str:
    return _bcrypt_lib.hashpw(password.encode(), _bcrypt_lib.gensalt()).decode()

def _verify_password(password: str, hashed: str) -> bool:
    return _bcrypt_lib.checkpw(password.encode(), hashed.encode())

# In-memory user store  {username: {"hashed_password": ..., "email": ...}}
# Replace with a real DB for production.
_USERS_FILE = "./data/users.json"

def _load_users() -> dict:
    try:
        if os.path.exists(_USERS_FILE):
            with open(_USERS_FILE) as f:
                return json.load(f)
    except Exception:
        pass
    return {}

def _save_users(users: dict):
    os.makedirs("data", exist_ok=True)
    with open(_USERS_FILE, "w") as f:
        json.dump(users, f, indent=2)

def _create_token(username: str) -> str:
    expire = datetime.datetime.utcnow() + datetime.timedelta(minutes=_TOKEN_EXPIRE)
    return jwt.encode({"sub": username, "exp": expire}, _SECRET_KEY, algorithm=_ALGORITHM)

def _decode_token(token: str) -> Optional[str]:
    try:
        payload = jwt.decode(token, _SECRET_KEY, algorithms=[_ALGORITHM])
        return payload.get("sub")
    except JWTError:
        return None

def get_current_user(creds: HTTPAuthorizationCredentials = Depends(_bearer)) -> str:
    if not creds:
        raise HTTPException(status_code=401, detail="Not authenticated")
    username = _decode_token(creds.credentials)
    if not username:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    users = _load_users()
    if username not in users:
        raise HTTPException(status_code=401, detail="User not found")
    return username

from agent import AIAgent
from config import setup_directories, settings

# ── App setup ────────────────────────────────────────────────────────────────
app = FastAPI(title="AI Agent API", version="1.0.0")

# Allow requests from local dev and any Vercel/production deployment
_ALLOWED_ORIGINS = [
    "http://localhost:3000",
    "http://localhost:8000",
    os.environ.get("FRONTEND_URL", ""),   # set this in Railway: FRONTEND_URL=https://nova-ui.vercel.app
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=[o for o in _ALLOWED_ORIGINS if o] or ["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

setup_directories()
os.makedirs("static", exist_ok=True)

# ── Agent instance (one per server, shared) ──────────────────────────────────
_agent: Optional[AIAgent] = None

def get_agent() -> AIAgent:
    global _agent
    if _agent is None:
        # Use the last known working model if saved
        model = _load_working_model()
        try:
            _agent = AIAgent(name="Nova", verbose=False, model=model)
        except Exception as e:
            raise HTTPException(status_code=503, detail=str(e))
    return _agent

# ── Persist the last working model across requests ───────────────────────────
_WORKING_MODEL_FILE = "./data/last_working_model.txt"

def _save_working_model(model: str):
    try:
        with open(_WORKING_MODEL_FILE, "w") as f:
            f.write(model)
    except Exception:
        pass

# Map every known model name → its provider
_MODEL_PROVIDER_MAP: dict[str, str] = {
    # Google
    "gemini-2.5-flash-lite": "google", "gemini-2.5-flash": "google",
    "gemini-2.0-flash": "google", "gemini-flash-latest": "google",
    "gemini-3-flash-preview": "google",
    # OpenAI
    "gpt-4o-mini": "openai", "gpt-4o": "openai", "gpt-3.5-turbo": "openai",
    # Anthropic
    "claude-3-5-haiku-20241022": "anthropic", "claude-3-5-sonnet-20241022": "anthropic",
    "claude-3-opus-20240229": "anthropic",
    # ICA
    "gpt-5.2": "ica",
    "meta-llama/llama-3-3-70b-instruct": "ica", "meta-llama/llama-3-1-8b-instruct": "ica",
    "ibm/granite-3-3-8b-instruct": "ica", "ibm/granite-3-3-2b-instruct": "ica",
    "mistralai/mistral-large": "ica",
}

_KEY_ATTR_MAP: dict[str, str] = {
    "google": "google_api_key",
    "openai": "openai_api_key",
    "anthropic": "anthropic_api_key",
    "ica": "ica_api_key",
}

def _load_working_model() -> Optional[str]:
    """
    Load the last saved model only if its provider's API key is still configured.
    If the key is missing, fall back to the default provider's model instead of
    crashing with a quota / auth error on startup.
    """
    try:
        if os.path.exists(_WORKING_MODEL_FILE):
            saved = open(_WORKING_MODEL_FILE).read().strip()
            if not saved:
                return None
            from config import settings as _s
            provider = _MODEL_PROVIDER_MAP.get(saved)
            key_attr  = _KEY_ATTR_MAP.get(provider or "")
            if key_attr and getattr(_s, key_attr, ""):
                return saved          # key exists → safe to use
            # Key missing — clear the stale file and fall back to default
            try:
                os.remove(_WORKING_MODEL_FILE)
            except Exception:
                pass
            return None               # caller will use the .env default
    except Exception:
        pass
    return None

# ── Request / Response models ─────────────────────────────────────────────────
class ChatRequest(BaseModel):
    message: str

class ChatResponse(BaseModel):
    reply: str
    turns: int

class StatusResponse(BaseModel):
    provider: str
    model: str
    tools: list
    turns: int

# ── Auth request/response models ─────────────────────────────────────────────
class SignupRequest(BaseModel):
    username: str
    password: str
    email: Optional[str] = None

class LoginRequest(BaseModel):
    username: str
    password: str

class AuthResponse(BaseModel):
    access_token: str
    username: str

# ── Auth routes ───────────────────────────────────────────────────────────────
@app.post("/api/auth/signup", response_model=AuthResponse)
def auth_signup(req: SignupRequest):
    import re
    if not re.match(r'^[a-zA-Z0-9_]{3,32}$', req.username):
        raise HTTPException(status_code=400, detail="Username must be 3–32 chars (letters, numbers, underscores)")
    if len(req.password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters")
    users = _load_users()
    if req.username in users:
        raise HTTPException(status_code=409, detail="Username already taken")
    users[req.username] = {
        "hashed_password": _hash_password(req.password),
        "email": req.email or "",
    }
    _save_users(users)
    return {"access_token": _create_token(req.username), "username": req.username}

@app.post("/api/auth/login", response_model=AuthResponse)
def auth_login(req: LoginRequest):
    users = _load_users()
    user  = users.get(req.username)
    if not user or not _verify_password(req.password, user["hashed_password"]):
        raise HTTPException(status_code=401, detail="Invalid username or password")
    return {"access_token": _create_token(req.username), "username": req.username}

@app.get("/api/auth/me")
def auth_me(username: str = Depends(get_current_user)):
    users = _load_users()
    user  = users.get(username, {})
    return {"username": username, "email": user.get("email", "")}

# ── Root redirect — UI is now served by Next.js on port 3000 ──────────────────
@app.get("/")
def root():
    from fastapi.responses import JSONResponse
    return JSONResponse({"message": "Nova AI API is running. Open http://localhost:3000 for the UI."})

@app.get("/api/status")
async def status():
    try:
        agent = get_agent()
        info = agent.get_info()
        return {
            "provider": info["provider"],
            "model":    agent.llm.model_name if hasattr(agent.llm, "model_name") else info["model"],
            "tools":    info["tools"],
            "turns":    info["memory_stats"]["turns"],
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=503, detail=str(e))

@app.post("/api/chat", response_model=ChatResponse)
async def chat(req: ChatRequest):
    """Non-streaming fallback — kept for compatibility."""
    if not req.message.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty")
    agent = get_agent()
    loop = asyncio.get_event_loop()
    reply = await loop.run_in_executor(None, agent.run, req.message.strip())
    current_model = agent.llm.model if hasattr(agent.llm, "model") else None
    if current_model:
        _save_working_model(current_model)
    return {"reply": reply, "turns": agent.memory.get_stats()["turns"]}


@app.post("/api/chat/stream")
async def chat_stream(req: ChatRequest):
    """
    Server-Sent Events streaming endpoint.
    Runs guardrails + cache synchronously, then streams LLM tokens.
    Each SSE event is JSON: {"type": "token"|"done"|"error", "data": "..."}
    """
    import json as _json
    from guardrails import get_guardrails, GuardrailViolation
    from cache import get_cache
    from tracer import get_tracer, EventType

    if not req.message.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty")

    msg = req.message.strip()
    agent = get_agent()

    async def event_generator():
        tracer     = get_tracer()
        guardrails = get_guardrails()
        cache      = get_cache()
        trace      = tracer.start(agent.memory.session_id, msg)

        def sse(event_type: str, data: str) -> str:
            return f"data: {_json.dumps({'type': event_type, 'data': data})}\n\n"

        try:
            # ── Guardrails (instant, no streaming needed) ──────────────────
            try:
                hitl_req = guardrails.check(msg)
                if hitl_req:
                    trace.log(EventType.HITL_PAUSE, {"rule_id": hitl_req.rule_id})
                    reply = (
                        f"This request requires human approval.\n\n"
                        f"**Reason:** {hitl_req.reason}"
                    )
                    yield sse("token", reply)
                    yield sse("done", str(agent.memory.get_stats()["turns"]))
                    tracer.finish(trace, reply)
                    return
            except GuardrailViolation as gv:
                trace.log(EventType.GUARDRAIL_DENY, {"rule_id": gv.rule_id})
                reply = f"Blocked: {gv.reason}"
                yield sse("token", reply)
                yield sse("done", str(agent.memory.get_stats()["turns"]))
                tracer.finish(trace, reply)
                return

            # ── Cache hit — inject as reference, let LLM rephrase ─────────
            cached = cache.get(msg)

            # ── LLM streaming ──────────────────────────────────────────────
            from rag import get_rag
            rag = get_rag()
            rag_context = rag.format_context(msg)
            from langchain_core.messages import HumanMessage, SystemMessage
            system = agent.system_message
            if rag_context:
                system = f"{agent.system_message}\n\n{rag_context}"
            if cached:
                trace.log(EventType.CACHE_HIT, {"preview": cached[:60]})
                system = (
                    f"{system}\n\n"
                    f"[Previous answer for reference — rephrase naturally, do not repeat verbatim]:\n{cached}"
                )

            messages = [SystemMessage(content=system)]
            messages.extend(agent.memory.get_messages())
            messages.append(HumanMessage(content=msg))

            trace.log(EventType.LLM_CALL, {"model": getattr(agent.llm, "model", "?")})

            full_reply = ""
            # Stream tokens via astream_events
            async for event in agent._graph.astream_events(
                {"messages": messages}, version="v2"
            ):
                kind = event.get("event", "")
                # Chat model streaming token
                if kind == "on_chat_model_stream":
                    chunk = event.get("data", {}).get("chunk")
                    if chunk:
                        token = ""
                        if hasattr(chunk, "content"):
                            c = chunk.content
                            if isinstance(c, str):
                                token = c
                            elif isinstance(c, list):
                                token = " ".join(
                                    p.get("text", "") for p in c if isinstance(p, dict)
                                )
                        if token:
                            full_reply += token
                            yield sse("token", token)
                # Tool call notification
                elif kind == "on_tool_start":
                    tool_name = event.get("name", "tool")
                    yield sse("tool", f"Using tool: {tool_name}...")

            # Persist to memory + cache
            if full_reply:
                agent.memory.add_message(msg, full_reply)
                cache.set(msg, full_reply)
                tracer.finish(trace, full_reply)
                model = getattr(agent.llm, "model", None)
                if model:
                    _save_working_model(model)

            yield sse("done", str(agent.memory.get_stats()["turns"]))

        except Exception as e:
            trace.log(EventType.ERROR, {"error": str(e)[:200]})
            tracer.finish(trace, f"Error: {e}")
            yield sse("error", str(e)[:300])
            yield sse("done", "0")

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )

# ── Model switcher ────────────────────────────────────────────────────────────
class SwitchModelRequest(BaseModel):
    provider: str
    model: str

AVAILABLE_MODELS = {
    "google": [
        "gemini-2.5-flash-lite",
        "gemini-2.5-flash",
        "gemini-2.0-flash",
        "gemini-flash-latest",
        "gemini-3-flash-preview",
    ],
    "openai": [
        "gpt-4o-mini",
        "gpt-4o",
        "gpt-3.5-turbo",
    ],
    "anthropic": [
        "claude-3-5-haiku-20241022",
        "claude-3-5-sonnet-20241022",
        "claude-3-opus-20240229",
    ],
    "ica": [
        "gpt-5.2",
        "meta-llama/llama-3-3-70b-instruct",
        "meta-llama/llama-3-1-8b-instruct",
        "ibm/granite-3-3-8b-instruct",
        "ibm/granite-3-3-2b-instruct",
        "mistralai/mistral-large",
    ],
}

@app.get("/api/models")
def list_models():
    from config import settings
    configured = {
        "openai":    bool(settings.openai_api_key),
        "anthropic": bool(settings.anthropic_api_key),
        "google":    bool(settings.google_api_key),
        "ica":       bool(settings.ica_api_key),
    }
    return {
        "current_provider": settings.default_llm_provider,
        "current_model": _load_working_model() or settings.google_model,
        "available": AVAILABLE_MODELS,
        "configured": configured,
    }

# Map provider → settings key name (for pre-flight check)
_PROVIDER_KEY_MAP = {
    "openai":    "openai_api_key",
    "anthropic": "anthropic_api_key",
    "google":    "google_api_key",
    "ica":       "ica_api_key",
}

@app.post("/api/models/switch")
def switch_model(req: SwitchModelRequest):
    global _agent
    if req.provider not in AVAILABLE_MODELS:
        raise HTTPException(status_code=400, detail=f"Unknown provider: {req.provider}")
    if req.model not in AVAILABLE_MODELS[req.provider]:
        raise HTTPException(status_code=400, detail=f"Unknown model: {req.model}")

    # Pre-flight: check the API key is configured before trying to build the agent
    from config import settings as _s
    key_attr = _PROVIDER_KEY_MAP.get(req.provider)
    if key_attr and not getattr(_s, key_attr, ""):
        raise HTTPException(
            status_code=400,
            detail=(
                f"No API key found for provider '{req.provider}'. "
                f"Add {key_attr.upper()} to your .env file and restart the server."
            ),
        )

    try:
        # Rebuild agent with new provider/model, keep existing memory
        old_memory = _agent.memory if _agent else None
        _agent = AIAgent(name="Nova", verbose=False, provider=req.provider, model=req.model)
        if old_memory:
            _agent.memory = old_memory   # preserve conversation history
        _save_working_model(req.model)
        return {"message": f"Switched to {req.provider} / {req.model}", "provider": req.provider, "model": req.model}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/clear")
def clear():
    agent = get_agent()
    agent.clear_memory()
    from cache import get_cache
    get_cache().clear()
    return {"message": "Memory and cache cleared"}

@app.get("/api/cache")
def cache_stats():
    from cache import get_cache
    return get_cache().stats()

@app.get("/api/history")
def history():
    agent = get_agent()
    msgs = agent.memory.get_messages()
    result = []
    for m in msgs:
        role = "user" if m.__class__.__name__ == "HumanMessage" else "assistant"
        result.append({"role": role, "content": m.content})
    return {"history": result}

# ── Tracing ───────────────────────────────────────────────────────────────────
@app.get("/api/traces")
def traces():
    from tracer import get_tracer
    return {"traces": get_tracer().get_recent(20)}

@app.get("/api/traces/stats")
def trace_stats():
    from tracer import get_tracer
    return get_tracer().get_stats()

# ── RAG ───────────────────────────────────────────────────────────────────────
class IngestRequest(BaseModel):
    text: str
    source: str = "manual"

@app.post("/api/rag/ingest")
def rag_ingest(req: IngestRequest):
    from rag import get_rag
    result = get_rag().ingest_text(req.text, req.source)
    return {"message": result}

@app.get("/api/rag/stats")
def rag_stats():
    from rag import get_rag
    return get_rag().get_stats()

# ── Guardrails & HITL ─────────────────────────────────────────────────────────
@app.get("/api/guardrails")
def guardrail_rules():
    from guardrails import get_guardrails
    return get_guardrails().get_all_rules()

@app.get("/api/hitl/pending")
def hitl_pending():
    from guardrails import get_guardrails
    return {"pending": get_guardrails().get_pending_hitl()}

class HITLResolve(BaseModel):
    rule_id: str
    approved: bool

@app.post("/api/hitl/resolve")
def hitl_resolve(req: HITLResolve):
    from guardrails import get_guardrails
    ok = get_guardrails().resolve_hitl(req.rule_id, req.approved)
    return {"resolved": ok, "approved": req.approved}

# ── Red Team ──────────────────────────────────────────────────────────────────
@app.post("/api/redteam/run")
async def redteam_run():
    """Run full red team security audit and return the report."""
    loop = asyncio.get_event_loop()
    from redteam import run_redteam
    report = await loop.run_in_executor(None, lambda: run_redteam(verbose=False))
    return report

# ── Agent Settings ────────────────────────────────────────────────────────────
_AGENT_SETTINGS_FILE = "./data/agent_settings.json"

_DEFAULT_SETTINGS = {
    "name":        "Nova",
    "temperature": 0.7,
    "max_tokens":  4096,
    "memory_type": "buffer",
    "verbose":     False,
    "system_message": "",
}

def _load_agent_settings() -> dict:
    try:
        if os.path.exists(_AGENT_SETTINGS_FILE):
            with open(_AGENT_SETTINGS_FILE) as f:
                return {**_DEFAULT_SETTINGS, **json.load(f)}
    except Exception:
        pass
    return dict(_DEFAULT_SETTINGS)

def _save_agent_settings(s: dict):
    os.makedirs("data", exist_ok=True)
    with open(_AGENT_SETTINGS_FILE, "w") as f:
        json.dump(s, f, indent=2)

class AgentSettingsRequest(BaseModel):
    name:           Optional[str]   = None
    temperature:    Optional[float] = None
    max_tokens:     Optional[int]   = None
    memory_type:    Optional[str]   = None
    verbose:        Optional[bool]  = None
    system_message: Optional[str]   = None

@app.get("/api/agent/settings")
def get_agent_settings():
    return _load_agent_settings()

@app.post("/api/agent/settings")
def update_agent_settings(req: AgentSettingsRequest):
    global _agent
    s = _load_agent_settings()
    if req.name           is not None: s["name"]           = req.name
    if req.temperature    is not None: s["temperature"]    = max(0.0, min(2.0, req.temperature))
    if req.max_tokens     is not None: s["max_tokens"]     = max(256, min(8192, req.max_tokens))
    if req.memory_type    is not None: s["memory_type"]    = req.memory_type
    if req.verbose        is not None: s["verbose"]        = req.verbose
    if req.system_message is not None: s["system_message"] = req.system_message
    _save_agent_settings(s)
    # Rebuild agent with new settings
    try:
        old_memory = _agent.memory if _agent else None
        _agent = AIAgent(
            name=s["name"],
            verbose=s["verbose"],
            model=_load_working_model(),
            temperature=s["temperature"],
            system_message=s["system_message"] or None,
        )
        if old_memory:
            _agent.memory = old_memory
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"ok": True, "settings": s}


# ── File Upload ───────────────────────────────────────────────────────────────
_UPLOAD_DIR = "./data/uploads"
os.makedirs(_UPLOAD_DIR, exist_ok=True)

_ALLOWED_EXTENSIONS = {
    # images
    ".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp",
    # documents
    ".pdf", ".txt", ".md", ".csv", ".json",
    ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx",
}
_MAX_UPLOAD_BYTES = 20 * 1024 * 1024  # 20 MB


@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...), username: str = Depends(get_current_user)):
    """Accept an uploaded file/image, save it, return its URL + extracted text."""
    import io, mimetypes

    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in _ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"File type '{ext}' not allowed.")

    content = await file.read()
    if len(content) > _MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large (max 20 MB).")

    # Save with a unique name
    safe_name = f"{secrets.token_hex(8)}{ext}"
    save_path  = os.path.join(_UPLOAD_DIR, safe_name)
    with open(save_path, "wb") as f:
        f.write(content)

    # Extract text content for the agent
    extracted = ""
    mime = (mimetypes.guess_type(file.filename or "")[0] or "")

    if mime.startswith("image/"):
        extracted = f"[Image uploaded: {file.filename}]"

    elif ext == ".txt" or ext == ".md" or ext == ".csv":
        try:
            extracted = content.decode("utf-8", errors="replace")[:8000]
        except Exception:
            pass

    elif ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            extracted = "\n".join(p.extract_text() or "" for p in reader.pages)[:8000]
        except Exception as e:
            extracted = f"[PDF could not be parsed: {e}]"

    elif ext in (".docx",):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            extracted = "\n".join(p.text for p in doc.paragraphs)[:8000]
        except Exception as e:
            extracted = f"[DOCX could not be parsed: {e}]"

    elif ext == ".json":
        try:
            extracted = json.dumps(json.loads(content), indent=2)[:8000]
        except Exception:
            extracted = content.decode("utf-8", errors="replace")[:8000]

    return {
        "filename": file.filename,
        "saved_as": safe_name,
        "url":      f"/api/uploads/{safe_name}",
        "mime":     mime,
        "is_image": mime.startswith("image/"),
        "text":     extracted,
    }


@app.get("/api/uploads/{filename}")
async def serve_upload(filename: str):
    """Serve a previously uploaded file."""
    import mimetypes
    path = os.path.join(_UPLOAD_DIR, filename)
    if not os.path.exists(path) or ".." in filename:
        raise HTTPException(status_code=404, detail="File not found")
    mime = mimetypes.guess_type(filename)[0] or "application/octet-stream"
    return FileResponse(path, media_type=mime)


# ── Export to DOCX ────────────────────────────────────────────────────────────
class ExportRequest(BaseModel):
    title:    str = "Nova AI Response"
    messages: List[dict]   # [{"role": "user"|"assistant", "content": "..."}]


@app.post("/api/export/docx")
def export_docx(req: ExportRequest):
    """Convert a list of chat messages to a .docx file and return it."""
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_para = doc.add_heading(req.title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Timestamp
    ts = doc.add_paragraph(f"Exported: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if ts.runs:
        ts.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x99)
        ts.runs[0].font.size = Pt(9)

    doc.add_paragraph()  # spacer

    for msg in req.messages:
        role    = msg.get("role", "")
        content = msg.get("content", "")

        # Role label
        label = doc.add_paragraph()
        run   = label.add_run("You:" if role == "user" else "Nova AI:")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = (
            RGBColor(0x3b, 0x82, 0xd4) if role == "user"
            else RGBColor(0x7c, 0x5c, 0xd8)
        )

        # Content — split on newlines for proper paragraphs
        for line in content.split("\n"):
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(11)

        doc.add_paragraph()  # spacer between messages

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = "".join(c for c in req.title if c.isalnum() or c in " _-")[:40].strip() or "export"
    filename   = f"{safe_title.replace(' ', '_')}.docx"

    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Entry point ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("\n" + "="*50)
    print("  AI Agent Web Interface")
    print("="*50)
    print("  Open: http://localhost:8000")
    print("  Stop: Ctrl+C")
    print("="*50 + "\n")
    uvicorn.run("server:app", host="0.0.0.0", port=8000, reload=False)
